"""
backend/routes/downloads.py
Download routes — CSV, PDF, Word audit report (per session).
"""
import io
from fastapi import APIRouter, Request
from fastapi.responses import RedirectResponse, StreamingResponse, HTMLResponse

from utils.logger  import get_logger
from backend.services import session_service

log    = get_logger("backend.routes.downloads")
router = APIRouter(prefix="/download")


@router.get("/csv")
async def download_csv(request: Request):
    _, session = session_service.get_or_create(request)
    state = session.get("state")
    if not state:
        return RedirectResponse("/")
    from src.extractors.validate_procurement import get_display_columns
    df = state.get("df_validation")
    if df is None or df.empty:
        return RedirectResponse("/")
    cols = [c for c in get_display_columns() if c in df.columns]
    buf  = io.StringIO()
    df[cols].to_csv(buf, index=True)
    return StreamingResponse(
        io.BytesIO(buf.getvalue().encode()),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=validation_report.csv"}
    )


@router.get("/pdf")
async def download_pdf(request: Request):
    _, session = session_service.get_or_create(request)
    state = session.get("state")
    if not state:
        return RedirectResponse("/")
    audit = state.get("audit_report", "")
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import mm

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=20*mm, rightMargin=20*mm,
                                topMargin=20*mm, bottomMargin=20*mm)
        styles  = getSampleStyleSheet()
        t_style = ParagraphStyle("t", parent=styles["Heading1"], fontSize=16, spaceAfter=12)
        b_style = ParagraphStyle("b", parent=styles["Normal"],   fontSize=10, leading=14, spaceAfter=6)
        story   = [Paragraph("Procurement Audit Report", t_style), Spacer(1, 6*mm)]
        for line in audit.split("\n"):
            line = line.strip()
            if not line:               story.append(Spacer(1, 3*mm))
            elif line.startswith("#"): story.append(Paragraph(line.lstrip("#").strip(), styles["Heading2"]))
            else:                      story.append(Paragraph(line, b_style))
        doc.build(story)
        buf.seek(0)
        return StreamingResponse(
            buf, media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=audit_report.pdf"}
        )
    except ImportError:
        return HTMLResponse("pip install reportlab", status_code=500)


@router.get("/word")
async def download_word(request: Request):
    _, session = session_service.get_or_create(request)
    state = session.get("state")
    if not state:
        return RedirectResponse("/")
    audit = state.get("audit_report", "")
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Procurement Audit Report", 0)
        for line in audit.split("\n"):
            line = line.strip()
            if not line:                  doc.add_paragraph("")
            elif line.startswith("####"): doc.add_heading(line.lstrip("#").strip(), level=4)
            elif line.startswith("###"):  doc.add_heading(line.lstrip("#").strip(), level=3)
            elif line.startswith("##"):   doc.add_heading(line.lstrip("#").strip(), level=2)
            elif line.startswith("#"):    doc.add_heading(line.lstrip("#").strip(), level=1)
            else:                         doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=audit_report.docx"}
        )
    except ImportError:
        return HTMLResponse("pip install python-docx", status_code=500)
